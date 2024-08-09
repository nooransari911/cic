from flask import Flask, jsonify, Request, Response, send_file, url_for, render_template_string
import flask
import functions_framework
from help import convert_request
from testim import get_images_from_drive, images
from docx import Document
from docx.shared import Inches, Pt, RGBColor
import io, os



app = Flask(__name__)

index_html = """
    <!DOCTYPE html>
    <html lang="en">
    <head>
        <meta charset="UTF-8">\
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>Image Heading Form</title>
    </head>

    <body>
        <h1>Enter Headings for Images</h1>

        <form method="POST">
            {% for image in images %}
                <div>
                    <br><br>
                    <img src={{image}} width="800">
                    <br><br>
                    <label for="{{image}}">Heading for {{ image }}:</label>
                    <br><br>
                    <input type="text" id="{{image}}" name="{{image}}" required>
                </div>
            {% endfor %}
            <br>
            <button type="submit">Generate Document</button>
        </form>
    </body>
    </html>
""".replace("\n", " ")

render_html = """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Document Generated</title>
</head>
<body>
    <h1>Document Generated Successfully</h1>
    <p>You can download the generated document <a href="{{ url_for('static', filename='generated_file.docx') }}">here</a>.</p>
</body>
</html>
""".replace("\n", " ")


'''
# json response
@app.route('/')
def hello_world():
  method = flask.request.method
  path = flask.request.path
  url = flask.request.url
  headers = dict (flask.request.headers)
  args = flask.request.args.to_dict()
  #data = flask.request.get_data()

  return jsonify (message=
    f"<h1>Hello, World!</h1>\
    <h1>Status code: 200</h1>\
    <h1>Request method: {method}</h1>\
    <h1>Request path: {path}</h1>\
    <h2>Request url: {url}</h2>\
    <h2>Request headers: {headers}</h2>\
    <h2>Request args: {args}</h2>",
    method=method,
    url=url,
    headers=headers,
    args=args)
    #data=data.decode ('utf-8'))
'''


@app.route('/co')
def hello_world():
  method = flask.request.method
  path = flask.request.path
  url = flask.request.url
  headers = dict (flask.request.headers)
  args = flask.request.args.to_dict()
  #data = flask.request.get_data()

  return Response (response=
    f"<h1>Hello, World!</h1>\
    <h1>Status code: 200</h1>\
    <h1>Request method: {method}</h1>\
    <h1>Request path: {path}</h1>\
    <h2>Request url: {url}</h2>\
    <h2>Request headers: {headers}</h2>\
    <h2>Request args: {args}</h2>",
    status=200,
    headers=headers)


@app.route ("/", methods=["GET", "POST"])
def compile_image ():
    # gdrive()
    # print ("return from refresh")
    folder_id = '1zSQOhdq3dv02PPyoqFapnZI0fx43jo_u'
    image_ids = get_images_from_drive(folder_id)
    image_list_rel = [url_for ('serve_image', image_id=image_id) for image_id in image_ids]

    #image_list_rel = [os.path.join ("/static/image/", im) for im in os.listdir(IMAGE_DIR) if im.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp'))]
    image_list_rel = sorted (image_list_rel)
    print (image_list_rel)
    #image_list_abs = [os.path.join(IMAGE_DIR, im) for im in os.listdir(IMAGE_DIR) if im.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp'))]


    doc = Document()
    section = doc.sections[0]
    margins = {
        'top': 0.4,  # 1 inch
        'bottom': 0.4,  # 1 inch
        'left': 0.4,  # 1 inch
        'right': 0.4  # 1 inch
    }
    section.top_margin = Inches(margins['top'])
    section.bottom_margin = Inches(margins['bottom'])
    section.left_margin = Inches(margins['left'])
    section.right_margin = Inches(margins['right'])


    if (flask.request.method=="GET"):
        return render_template_string (index_html, images=image_list_rel)

    if (flask.request.method=="POST"):
        for im in image_list_rel:

            para = doc.add_heading(level=1)
            run = para.add_run (request.form.get (im, ""))
            run.font.size = Pt (20)
            run.font.color.rgb = RGBColor (0, 0, 0)
            run.font.bold = True

            doc.add_picture (PATH_APP[:-1]+im, width=Inches (3.0))


        doc_io = io.BytesIO()
        doc.save (doc_io)
        doc_io.seek (0)  # Go back to the beginning of the BytesIO buffer

        response = make_response(send_file(doc_io, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',\
                                            as_attachment=True, attachment_filename='generated.docx'))
        response.headers['Content-Disposition'] = 'attachment; generated.docx'
        return response


@app.route('/refresh')
def index():
    folder_id = '1zSQOhdq3dv02PPyoqFapnZI0fx43jo_u'
    image_ids = get_images_from_drive(folder_id)
    image_urls = [url_for ('serve_image', image_id=image_id) for image_id in image_ids]
    print (image_urls)
    return render_template_string (index_html, image_urls=image_urls)


@app.route('/image/<image_id>')
def serve_image(image_id):
    image_data = images.get(image_id)
    if image_data:
        return send_file(io.BytesIO(image_data), mimetype='image/jpeg')
    else:
        return "Image not found", 404





# Entry point for Google Cloud Functions
@functions_framework.http
def hello_world (request):
  environ_dictionary = convert_request (request)

  with app.test_request_context (environ_base=environ_dictionary.environ):
        fresponse = app.full_dispatch_request ()

        #return fresponse
        
        return Response (
            response=fresponse.get_data(),
            status=fresponse.status_code,
            headers=dict(fresponse.headers))
        


if __name__ == '__main__':
    app.run(debug=True)
